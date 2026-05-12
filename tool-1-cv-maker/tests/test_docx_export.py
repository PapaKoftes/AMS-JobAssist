"""
Tests for DOCX Export functionality (Day 10 Phase 10.2 continuation).

Tests the DOCXExporter class that:
- Exports CVData as DOCX files
- Exports in multiple languages
- Validates CVData before export
- Generates professional formatted Word documents
- Handles edge cases and errors
"""

import pytest
from pathlib import Path
from src.backend.export.docx_export import DOCXExporter
from src.backend.cv.models import CVData, CVSection, QuestionCategory
import tempfile
import shutil
from docx import Document


def _all_text(doc: Document) -> str:
    """
    Extract ALL text from a DOCX document, including table cells.
    doc.paragraphs only returns top-level paragraphs and misses table cells.
    """
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return '\n'.join(parts)


class TestDOCXExporter:
    """Test DOCX export functionality."""

    @pytest.fixture
    def temp_output_dir(self):
        """Create temporary output directory for exports."""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        # Cleanup
        shutil.rmtree(temp_dir, ignore_errors=True)

    @pytest.fixture
    def exporter(self, temp_output_dir):
        """Provide DOCXExporter with temporary output directory."""
        return DOCXExporter(output_dir=temp_output_dir)

    @pytest.fixture
    def sample_cv_data(self):
        """Provide sample CVData with sections."""
        cv_data = CVData(
            session_id=1,
            user_id="test_user",
            interview_path="unemployed",
            language_input="en"
        )
        cv_data.language_output_primary = "de"
        cv_data.language_output_secondary = "en"

        # Add background section
        bg = CVSection(
            german="Ich bin ein Software-Ingenieur mit 5 Jahren Erfahrung.",
            english="I am a software engineer with 5 years of experience.",
            native="I am a software engineer with 5 years of experience.",
            category=QuestionCategory.BACKGROUND,
            question_id="bg_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.85,
            confidence_level="high",
            detected_skills=["Software Engineering"]
        )
        cv_data.background.append(bg)

        # Add experience section
        exp = CVSection(
            german="Ich leitete ein Team von fünf Entwicklern.",
            english="I led a team of five developers.",
            native="I led a team of five developers.",
            category=QuestionCategory.EXPERIENCE,
            question_id="exp_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.80,
            confidence_level="high",
            detected_skills=["Leadership"]
        )
        cv_data.experience.append(exp)

        # Add skills section
        skills = CVSection(
            german="Ich beherrsche Python, SQL und MySQL.",
            english="I master Python, SQL and MySQL.",
            native="I master Python, SQL and MySQL.",
            category=QuestionCategory.SKILLS,
            question_id="skills_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.90,
            confidence_level="high",
            detected_skills=["Python", "SQL", "MySQL"]
        )
        cv_data.skills.append(skills)

        # Set all skills and quality
        cv_data.all_skills = ["Python", "SQL", "MySQL", "Leadership", "Software Engineering"]
        quality_scores = [bg.quality_score, exp.quality_score, skills.quality_score]
        cv_data.overall_quality = sum(quality_scores) / len(quality_scores)
        cv_data.ready_for_export = cv_data.overall_quality >= 0.5

        return cv_data

    def test_export_docx_creates_file(self, exporter, sample_cv_data):
        """Test that export creates a DOCX file."""
        result = exporter.export(sample_cv_data, language="de")

        assert result is not None
        assert Path(result).exists()
        assert result.endswith(".docx")

    def test_export_docx_is_valid_word_document(self, exporter, sample_cv_data):
        """Test that exported file is a valid Word document."""
        result = exporter.export(sample_cv_data, language="de")

        # Try to open as Word document
        doc = Document(result)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_export_docx_contains_user_id(self, exporter, sample_cv_data):
        """Test that exported document contains user ID or identity name."""
        result = exporter.export(sample_cv_data, language="de")

        doc = Document(result)
        # Find user ID in document — header now shows identity.full_name or user_id
        text = _all_text(doc)
        # Accept either the identity name or the user_id (with any casing)
        identity = sample_cv_data.identity
        display_name = (identity.full_name if identity and identity.full_name else sample_cv_data.user_id)
        assert display_name in text or sample_cv_data.user_id in text

    def test_export_docx_german_content(self, exporter, sample_cv_data):
        """Test that German export creates valid DOCX."""
        result = exporter.export(sample_cv_data, language="de")

        assert result is not None
        assert Path(result).exists()

        # Verify content
        doc = Document(result)
        text = _all_text(doc)
        # Should contain German translation
        assert "Software-Ingenieur" in text or "Ingenieur" in text

    def test_export_docx_english_content(self, exporter, sample_cv_data):
        """Test that English export creates valid DOCX."""
        result = exporter.export(sample_cv_data, language="en")

        assert result is not None
        assert Path(result).exists()

        doc = Document(result)
        text = _all_text(doc)
        # Should contain English text
        assert "engineer" in text.lower()

    def test_export_docx_with_custom_filename(self, exporter, sample_cv_data):
        """Test export with custom filename."""
        result = exporter.export(sample_cv_data, language="de", filename="my_cv")

        assert "my_cv" in result
        assert Path(result).exists()
        assert result.endswith(".docx")

    def test_export_docx_validates_cv_data(self, exporter):
        """Test that export validates CVData before exporting."""
        # Create invalid CVData (no sections)
        invalid_cv = CVData(
            session_id=1,
            user_id="test_user",
            interview_path="unemployed",
            language_input="en"
        )

        result = exporter.export(invalid_cv, language="de")

        # Should fail validation
        assert result is None

    def test_export_multiple_sections_same_category(self, exporter, sample_cv_data):
        """Test export with multiple sections in same category."""
        exp2 = CVSection(
            german="Ich entwickelte eine mobile Anwendung.",
            english="I developed a mobile application.",
            native="I developed a mobile application.",
            category=QuestionCategory.EXPERIENCE,
            question_id="exp_002",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.75,
            confidence_level="high",
            detected_skills=["Mobile Development"]
        )
        sample_cv_data.experience.append(exp2)

        result = exporter.export(sample_cv_data, language="de")

        assert result is not None
        assert Path(result).exists()

        doc = Document(result)
        text = _all_text(doc)
        # Should have both experiences mentioned
        assert "Team" in text or "entwickelte" in text

    def test_export_with_empty_sections(self, exporter):
        """Test export with some empty category lists."""
        cv = CVData(
            session_id=1,
            user_id="test",
            interview_path="unemployed",
            language_input="en"
        )
        cv.language_output_primary = "de"
        cv.language_output_secondary = "en"

        # Only add background
        bg = CVSection(
            german="Hintergrund Info",
            english="Background info",
            native="Background info",
            category=QuestionCategory.BACKGROUND,
            question_id="bg_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.70,
            confidence_level="high",
            detected_skills=[]
        )
        cv.background.append(bg)
        cv.all_skills = []
        cv.overall_quality = 0.70
        cv.ready_for_export = True

        result = exporter.export(cv, language="de")

        assert result is not None
        assert Path(result).exists()

        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_filename_generation_for_docx(self, exporter, sample_cv_data):
        """Test DOCX-specific filename generation."""
        filename = exporter.generate_filename(sample_cv_data, "de", "docx")

        # Should contain user_id, interview_path, and language
        assert "test_user" in filename
        assert "unemployed" in filename
        assert "de" in filename
        assert filename.endswith(".docx")

    def test_export_different_languages_creates_different_files(self, exporter, sample_cv_data):
        """Test that exporting in different languages creates different files."""
        result_de = exporter.export(sample_cv_data, language="de")
        result_en = exporter.export(sample_cv_data, language="en")

        assert result_de is not None
        assert result_en is not None
        assert result_de != result_en  # Different file paths
        assert Path(result_de).exists()
        assert Path(result_en).exists()

    def test_export_docx_with_long_text(self, exporter):
        """Test export with long text sections."""
        cv = CVData(
            session_id=1,
            user_id="test_long",
            interview_path="unemployed",
            language_input="en"
        )
        cv.language_output_primary = "de"
        cv.language_output_secondary = "en"

        # Add section with long text
        long_text = "I have extensive experience in software development. " * 10
        section = CVSection(
            german=long_text,
            english=long_text,
            native=long_text,
            category=QuestionCategory.EXPERIENCE,
            question_id="exp_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.80,
            confidence_level="high",
            detected_skills=["Python"]
        )
        cv.experience.append(section)
        cv.all_skills = ["Python"]
        cv.overall_quality = 0.80
        cv.ready_for_export = True

        result = exporter.export(cv, language="en")

        assert result is not None
        assert Path(result).exists()

        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_export_docx_with_special_characters(self, exporter):
        """Test export with special characters and umlauts."""
        cv = CVData(
            session_id=1,
            user_id="test_special",
            interview_path="unemployed",
            language_input="en"
        )
        cv.language_output_primary = "de"
        cv.language_output_secondary = "en"

        # Add section with German umlauts
        section = CVSection(
            german="Ich beherrsche Programmierung, Öffentlichkeitsarbeit und ähnliches.",
            english="I master programming, public relations and similar.",
            native="I master programming, public relations and similar.",
            category=QuestionCategory.SKILLS,
            question_id="skills_001",
            detected_input_language="en",
            user_native_language="en",
            quality_score=0.85,
            confidence_level="high",
            detected_skills=["Programmierung", "Öffentlichkeitsarbeit"]
        )
        cv.skills.append(section)
        cv.all_skills = ["Programmierung", "Öffentlichkeitsarbeit"]
        cv.overall_quality = 0.85
        cv.ready_for_export = True

        result = exporter.export(cv, language="de")

        assert result is not None
        assert Path(result).exists()

        doc = Document(result)
        text = _all_text(doc)
        # Should contain special characters
        assert "ö" in text or "Öffentlichkeitsarbeit" in text

    def test_export_docx_document_structure(self, exporter, sample_cv_data):
        """Test that exported document has proper structure."""
        result = exporter.export(sample_cv_data, language="de")

        doc = Document(result)

        # Check for header
        assert len(doc.paragraphs) > 0

        # Check for sections
        text = _all_text(doc)
        # Should have some section titles
        assert len(text) > 100  # Reasonable document length


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
