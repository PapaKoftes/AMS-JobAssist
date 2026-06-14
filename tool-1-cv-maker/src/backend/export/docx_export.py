"""
DOCX Exporter - Export CVData as Austrian Tabellarischer Lebenslauf (Word format).

Generates .docx files following Austrian CV conventions:
  - Header table: name + contact details on the left, photo on the right
  - Section headings: bold, uppercase, with blue bottom border
  - Work/education entries: date | title / employer / bullets (tab-stop layout)
  - Font: Calibri 11pt (full umlaut support in Word)
  - Margins: 2 cm all sides
"""

import base64
import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from cv.models import CVData, CVSection
from cv.language_levels import display_label as _level_display_label
from export.base import CVExporter

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Translation tables
# ---------------------------------------------------------------------------
_LABELS_DE = {
    "background": "AUSBILDUNG",
    "experience": "BERUFSERFAHRUNG",
    "skills":     "KENNTNISSE",
    "motivation": "MOTIVATION",
    "training":   "WEITERBILDUNG & ZERTIFIKATE",
    "projects":   "PROJEKTE",
    "languages":  "SPRACHEN",
}
_LABELS_EN = {
    "background": "EDUCATION",
    "experience": "PROFESSIONAL EXPERIENCE",
    "skills":     "SKILLS",
    "motivation": "MOTIVATION",
    "training":   "TRAINING & CERTIFICATIONS",
    "projects":   "PROJECTS",
    "languages":  "LANGUAGES",
}

# Language-level display labels live in cv.language_levels (single source of truth).

# Photo dimensions (Austrian standard)
PHOTO_W_CM = 3.5
PHOTO_H_CM = 4.5

# Colours
COLOR_SECTION = RGBColor(26,  26,  26)   # near-black
COLOR_TEXT    = RGBColor(51,  51,  51)
COLOR_MUTED   = RGBColor(102, 102, 102)
COLOR_ACCENT  = RGBColor(44,  62,  80)   # dark blue


class DOCXExporter(CVExporter):
    """Export CVData as Austrian Tabellarischer Lebenslauf DOCX."""

    # ====================================================================
    # Public API
    # ====================================================================

    def export(
        self,
        cv_data: CVData,
        language: str = "de",
        filename: Optional[str] = None,
    ) -> Optional[str]:
        """
        Export CVData as DOCX file.

        Returns path to generated file, or None on failure.
        """
        try:
            logger.info(f"Exporting CVData as DOCX (language={language})")

            if not self.validate_cv_data(cv_data):
                logger.error("CVData validation failed")
                return None

            content = self.get_cv_content_for_language(cv_data, language)
            if content is None:
                logger.error("Failed to extract CV content")
                return None

            docx_filename = self.generate_filename(cv_data, language, "docx", filename)
            docx_path = self.output_dir / docx_filename

            doc = Document()
            self._setup_page_margins(doc)
            self._setup_default_style(doc)
            self._build_document(doc, cv_data, content, language)
            doc.save(str(docx_path))

            logger.info(f"DOCX export successful: {docx_path}")
            return str(docx_path)

        except Exception as e:
            logger.error(f"DOCX export error: {e}", exc_info=True)
            return None

    # ====================================================================
    # Document builder
    # ====================================================================

    def _build_document(
        self,
        doc: Document,
        cv_data: CVData,
        content: Dict[str, Any],
        language: str,
    ):
        """Populate the Word document with all CV content."""

        # 1. Header (name + contact + photo)
        self._add_header(doc, cv_data, language)

        # 2. Experience
        if cv_data.experience:
            self._add_section_heading(doc, "experience", language)
            groups = self._group_experience_sections(cv_data.experience)
            for block in groups:
                self._add_experience_block(doc, block, language)

        # 3. Education / Background
        if cv_data.background:
            self._add_section_heading(doc, "background", language)
            groups = self._group_experience_sections(cv_data.background)
            for block in groups:
                self._add_experience_block(doc, block, language)

        # 4. Training
        if cv_data.training:
            self._add_section_heading(doc, "training", language)
            groups = self._group_experience_sections(cv_data.training)
            for block in groups:
                self._add_experience_block(doc, block, language)

        # 5. Skills
        all_skills = content.get("all_skills", [])
        if all_skills or cv_data.skills:
            self._add_section_heading(doc, "skills", language)
            self._add_skills_block(doc, all_skills, cv_data.skills, language)

        # 6. Languages
        if cv_data.languages:
            self._add_section_heading(doc, "languages", language)
            self._add_languages_block(doc, cv_data.languages, language)

        # 7. Free-text sections
        for key in ("motivation", "projects"):
            section_list = getattr(cv_data, key, [])
            if section_list:
                self._add_section_heading(doc, key, language)
                for s in section_list:
                    if s.hidden:
                        continue
                    text = s.german if language == "de" else s.english
                    if text:
                        p = doc.add_paragraph(text)
                        p.paragraph_format.space_after = Pt(4)

        # 8. Austrian signature block (Ort, Datum + Unterschrift)
        self._add_signature_block(doc, cv_data, language)

        # 9. Footer
        self._add_footer(doc, cv_data, language)

    # ====================================================================
    # Header
    # ====================================================================

    def _add_header(self, doc: Document, cv_data: CVData, language: str):
        """
        Add a two-column header table:
          left:  name (large bold) + contact info
          right: photo (3.5 × 4.5 cm)
        """
        identity = cv_data.identity
        display_name = (
            identity.full_name
            if identity and identity.full_name
            else cv_data.user_id
        )

        # Total usable width ≈ 17 cm (A4 − 2×2 cm margins)
        photo_col_w = Cm(PHOTO_W_CM + 0.3)
        left_col_w  = Cm(17.0) - photo_col_w

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _remove_table_borders(tbl)

        left_cell  = tbl.cell(0, 0)
        right_cell = tbl.cell(0, 1)

        # Set column widths
        left_cell.width  = left_col_w
        right_cell.width = photo_col_w

        # ---- Left cell: name + contact -----------------------------
        # Name paragraph
        name_para = left_cell.paragraphs[0]
        name_run = name_para.add_run(display_name)
        name_run.font.size    = Pt(20)
        name_run.font.bold    = True
        name_run.font.color.rgb = COLOR_SECTION
        name_para.paragraph_format.space_after = Pt(4)

        # Contact lines
        if identity:
            contact_parts = []
            if identity.location:
                contact_parts.append(f"📍 {identity.location}")
            if identity.contact_phone:
                contact_parts.append(f"📞 {identity.contact_phone}")
            if identity.contact_email:
                contact_parts.append(f"✉ {identity.contact_email}")
            if contact_parts:
                cp = left_cell.add_paragraph("  ".join(contact_parts))
                cp.runs[0].font.size = Pt(9)
                cp.runs[0].font.color.rgb = COLOR_MUTED
                cp.paragraph_format.space_after = Pt(2)

            extra = []
            if identity.date_of_birth:
                extra.append(f"📅 Geburtsdatum: {identity.date_of_birth}")
            if identity.nationality:
                extra.append(f"🌍 Staatsangehörigkeit: {identity.nationality}")
            if extra:
                ep = left_cell.add_paragraph("  ".join(extra))
                ep.runs[0].font.size = Pt(9)
                ep.runs[0].font.color.rgb = COLOR_MUTED
                ep.paragraph_format.space_after = Pt(2)

        # ---- Right cell: photo ------------------------------------
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_bytes = None
        if identity and identity.photo:
            photo_bytes = self._decode_photo(identity.photo)
        if photo_bytes:
            try:
                img_io = io.BytesIO(photo_bytes)
                right_cell.paragraphs[0].add_run().add_picture(
                    img_io,
                    width=Cm(PHOTO_W_CM),
                    height=Cm(PHOTO_H_CM),
                )
            except Exception as exc:
                logger.warning(f"Could not embed photo: {exc}")

        # ---- Thick rule under header --------------------------------
        doc.add_paragraph()  # small gap
        self._add_horizontal_rule(doc, thickness="18", color="2C3E50")

    # ====================================================================
    # Section heading
    # ====================================================================

    def _add_section_heading(self, doc: Document, key: str, language: str):
        """Bold uppercase heading with blue bottom border."""
        label = (
            _LABELS_DE.get(key, key.upper())
            if language == "de"
            else _LABELS_EN.get(key, key.upper())
        )
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after  = Pt(2)
        run = heading.add_run(label)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = COLOR_SECTION
        self._set_paragraph_bottom_border(heading, color="AAAAAA")

    # ====================================================================
    # Experience / Education blocks
    # ====================================================================

    def _group_experience_sections(
        self, sections: List[CVSection]
    ) -> List[Dict]:
        """
        Group CVSections by base question_id.

        Sections ending in ``_employer``, ``_title``, ``_dates`` are merged
        with their parent entry.  Returns list of group dicts.
        """
        groups: Dict[str, Dict] = {}
        order: List[str] = []

        for s in sections:
            if s.hidden:
                continue
            qid = s.question_id or ""

            if qid.endswith("_employer"):
                base = qid[: -len("_employer")]
                groups.setdefault(base, {})["employer"] = s.german
                if base not in order:
                    order.append(base)
            elif qid.endswith("_title"):
                base = qid[: -len("_title")]
                groups.setdefault(base, {})["title"] = s.german
                if base not in order:
                    order.append(base)
            elif qid.endswith("_dates"):
                base = qid[: -len("_dates")]
                groups.setdefault(base, {})["dates"] = s.german
                if base not in order:
                    order.append(base)
            else:
                groups.setdefault(qid, {})["main"] = s
                if qid not in order:
                    order.append(qid)

        result = []
        for b in order:
            g = groups.get(b, {})
            main = g.get("main")
            # Prefer values folded onto the main section by the builder; fall
            # back to legacy separate _title/_employer/_dates sub-sections.
            title = g.get("title") or (getattr(main, "title", "") if main else "")
            employer = g.get("employer") or (getattr(main, "employer", "") if main else "")
            dates = g.get("dates")
            if not dates and main and getattr(main, "period", None):
                fmt = getattr(self, "_format_period", None)
                dates = fmt(main.period) if callable(fmt) else None
            result.append({
                "main":     main,
                "employer": employer or None,
                "title":    title or None,
                "dates":    dates,
            })
        return result

    def _add_experience_block(
        self, doc: Document, block: Dict, language: str
    ):
        """
        Add a single work/education entry as a two-column table row:
            DATE COL  |  TITLE (bold)
                      |  Employer (italic, muted)
                      |  • bullet
                      |  • bullet
        """
        main: Optional[CVSection] = block.get("main")

        _dates_raw = block.get("dates")
        dates_str = (_dates_raw.german or _dates_raw.english or "") if hasattr(_dates_raw, "german") else str(_dates_raw or "")
        if not dates_str and main and main.period:
            dates_str = self._format_period(main.period)
        dates_str = dates_str or "–"

        _title_raw    = block.get("title")
        _employer_raw = block.get("employer")
        title_str    = (_title_raw.german    if hasattr(_title_raw,    "german") else str(_title_raw    or ""))
        employer_str = (_employer_raw.german if hasattr(_employer_raw, "german") else str(_employer_raw or ""))

        if not title_str and main:
            title_str = main.german or ""

        bullets: List[str] = []
        if main and main.bullets:
            bullets = main.bullets

        # Two-column table: date (3.8 cm) | content
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _remove_table_borders(tbl)

        date_cell    = tbl.cell(0, 0)
        content_cell = tbl.cell(0, 1)
        date_cell.width    = Cm(3.8)
        content_cell.width = Cm(17.0 - 3.8)

        # Date cell
        dp = date_cell.paragraphs[0]
        dr = dp.add_run(dates_str)
        dr.font.size = Pt(9)
        dr.font.color.rgb = COLOR_MUTED

        # Content cell — title
        first_para = content_cell.paragraphs[0]
        if title_str:
            tr = first_para.add_run(title_str)
            tr.font.bold  = True
            tr.font.size  = Pt(10)
            tr.font.color.rgb = COLOR_TEXT
            first_para.paragraph_format.space_after = Pt(1)
        elif not title_str and not employer_str and not bullets and main:
            # Fallback: plain text
            text = main.german if language == "de" else main.english
            first_para.add_run(text or "").font.size = Pt(9)

        # Employer
        if employer_str:
            ep = content_cell.add_paragraph()
            er = ep.add_run(employer_str)
            er.font.italic = True
            er.font.size   = Pt(9)
            er.font.color.rgb = COLOR_MUTED
            ep.paragraph_format.space_after = Pt(2)

        # Bullets
        for bullet in bullets:
            bp = content_cell.add_paragraph()
            bt = bullet if bullet.startswith("•") else f"• {bullet}"
            br = bp.add_run(bt)
            br.font.size = Pt(9)
            br.font.color.rgb = COLOR_TEXT
            bp.paragraph_format.left_indent  = Pt(6)
            bp.paragraph_format.space_after  = Pt(1)

        # Gap after block
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ====================================================================
    # Skills
    # ====================================================================

    def _add_skills_block(
        self,
        doc: Document,
        all_skills: List[str],
        skills_sections: List[CVSection],
        language: str,
    ):
        items = all_skills if all_skills else [
            (s.german if language == "de" else s.english)
            for s in skills_sections
            if not s.hidden and (s.german or s.english)
        ]
        for item in items:
            p = doc.add_paragraph()
            r = p.add_run(f"• {item}")
            r.font.size = Pt(9)
            p.paragraph_format.left_indent  = Pt(6)
            p.paragraph_format.space_after  = Pt(2)

    # ====================================================================
    # Languages
    # ====================================================================

    def _add_languages_block(
        self, doc: Document, languages: List[Dict], language: str
    ):
        """Two-column table: language name | CEFR level."""
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _remove_table_borders(tbl)

        for lang_entry in languages:
            row = tbl.add_row()
            lang_name  = lang_entry.get("language", "")
            level_label = _level_display_label(lang_entry.get("level", ""))

            name_cell  = row.cells[0]
            level_cell = row.cells[1]
            name_cell.width  = Cm(8.5)
            level_cell.width = Cm(8.5)

            np = name_cell.paragraphs[0]
            nr = np.add_run(lang_name)
            nr.font.bold = True
            nr.font.size = Pt(9)

            lp = level_cell.paragraphs[0]
            lr = lp.add_run(level_label)
            lr.font.size = Pt(9)
            lr.font.color.rgb = COLOR_MUTED

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ====================================================================
    # Footer
    # ====================================================================

    def _add_signature_block(self, doc: Document, cv_data: CVData, language: str):
        """Austrian CV convention: place + date, blank signature space, typed name."""
        ident = getattr(cv_data, "identity", None)
        name = (getattr(ident, "full_name", "") or "").strip() if ident else ""
        if not name:
            return
        city = (getattr(ident, "location", "") or "").strip() if ident else ""
        date_str = datetime.now().strftime("%d.%m.%Y")
        ort_datum = f"{city}, {date_str}" if city else date_str
        doc.add_paragraph()  # spacing
        p1 = doc.add_paragraph(ort_datum)
        p1.runs[0].font.size = Pt(9)
        doc.add_paragraph()  # blank space for the handwritten signature
        p2 = doc.add_paragraph("______________________________")
        p2.runs[0].font.size = Pt(9)
        p3 = doc.add_paragraph(name)
        p3.runs[0].font.size = Pt(9)

    def _add_footer(self, doc: Document, cv_data: CVData, language: str):
        self._add_horizontal_rule(doc, thickness="6", color="CCCCCC")
        date_str    = datetime.now().strftime("%d.%m.%Y")
        quality_pct = f"{cv_data.overall_quality:.0%}"
        if language == "de":
            text = f"Lebenslauf erstellt: {date_str} | Qualität: {quality_pct} | AMS JobAssist"
        else:
            text = f"CV generated: {date_str} | Quality: {quality_pct} | AMS JobAssist"
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(7)
        p.runs[0].font.color.rgb = RGBColor(153, 153, 153)

    # ====================================================================
    # Helpers
    # ====================================================================

    def _decode_photo(self, photo_str: str) -> Optional[bytes]:
        """Decode photo from base64 data URI or file path."""
        if not photo_str:
            return None
        if photo_str.startswith("data:"):
            try:
                _header, b64data = photo_str.split(",", 1)
                return base64.b64decode(b64data)
            except Exception as exc:
                logger.warning(f"Base64 decode failed: {exc}")
                return None
        path = Path(photo_str)
        if path.exists():
            return path.read_bytes()
        logger.warning(f"Photo path not found: {photo_str}")
        return None

    def _format_period(self, period: Optional[Dict]) -> str:
        if not period:
            return ""
        start = self._format_month(period.get("start"))
        end   = self._format_month(period.get("end"))
        if start and end:
            return f"{start} – {end}"
        if start:
            return f"{start} – heute"
        return ""

    @staticmethod
    def _format_month(ym: Optional[str]) -> str:
        if not ym:
            return ""
        _months = ["", "Jan", "Feb", "Mär", "Apr", "Mai", "Jun",
                   "Jul", "Aug", "Sep", "Okt", "Nov", "Dez"]
        parts = ym.split("-")
        if len(parts) == 2:
            try:
                year, month = int(parts[0]), int(parts[1])
                if 1 <= month <= 12:
                    return f"{_months[month]} {year}"
            except ValueError:
                pass
        return ym

    def _setup_page_margins(self, doc: Document):
        """Set 2 cm margins on all sides."""
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2)
            section.right_margin  = Cm(2)

    def _setup_default_style(self, doc: Document):
        """Set default font to Calibri 10pt."""
        style = doc.styles["Normal"]
        font  = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = COLOR_TEXT

    def _add_horizontal_rule(
        self, doc: Document, thickness: str = "6", color: str = "CCCCCC"
    ):
        """Add a paragraph styled with a bottom border to simulate a <hr>."""
        p = doc.add_paragraph()
        self._set_paragraph_bottom_border(p, thickness=thickness, color=color)
        p.paragraph_format.space_after = Pt(4)

    @staticmethod
    def _set_paragraph_bottom_border(
        paragraph, thickness: str = "6", color: str = "AAAAAA"
    ):
        """Attach an OOXML bottom-border to a paragraph element."""
        pPr  = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    thickness)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Utility: strip table borders (set all to "none")
# ---------------------------------------------------------------------------

def _remove_table_borders(tbl):
    """Remove all borders from a table by setting them to 'none' via XML."""
    tbl_element = tbl._tbl
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_element.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"),  "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Quick smoke-test
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    from cv.models import CVData, CVSection, CVIdentity, QuestionCategory

    cv = CVData(
        session_id="1",
        user_id="test_user",
        interview_path="unemployed",
        language_input="de",
        language_output_primary="de",
        language_output_secondary="en",
    )
    cv.identity = CVIdentity(
        full_name="Ayşe Yılmaz",
        date_of_birth="15.03.1990",
        nationality="Türkisch",
        location="Wien, Österreich",
        contact_email="ayse@example.com",
        contact_phone="+43 676 1234567",
    )

    exp = CVSection(
        german="Lagermitarbeiterin",
        english="Warehouse Worker",
        native="Lagermitarbeiterin",
        category=QuestionCategory.EXPERIENCE,
        question_id="exp_001",
        detected_input_language="de",
        user_native_language="tr",
        quality_score=0.85,
        confidence_level="high",
        detected_skills=["Lagerlogistik"],
        period={"start": "2020-01", "end": "2022-03"},
        bullets=[
            "Übernahm Verantwortung für Lagerlogistik",
            "Koordinierte Wareneingang und -ausgang",
        ],
    )
    cv.experience.append(exp)

    cv.all_skills = ["Microsoft Office (Word, Excel, Outlook)", "Führerschein Klasse B"]
    cv.languages = [
        {"language": "Deutsch",  "code": "de", "level": "b2"},
        {"language": "Türkisch", "code": "tr", "level": "native"},
        {"language": "Englisch", "code": "en", "level": "a2"},
    ]
    cv.overall_quality = 0.85
    cv.ready_for_export = True

    exporter = DOCXExporter()
    path = exporter.export(cv, language="de", filename="test_austrian_cv")
    print(f"Exported to: {path}")
